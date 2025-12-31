import base64
import logging
from odoo import models, fields, api
from odoo.exceptions import UserError

_logger = logging.getLogger(__name__)


class DocumentsDocument(models.Model):
    _inherit = "documents.document"

    is_text_editable = fields.Boolean(
        "Is Text Editable", compute="_compute_is_text_editable", store=False
    )

    text_content = fields.Text(
        "Text Content", compute="_compute_text_content", inverse="_inverse_text_content"
    )

    @api.depends("mimetype")
    def _compute_is_text_editable(self):
        """Check if document can be edited as text"""
        editable_types = [
            "text/plain",
            "text/html",
            "text/css",
            "text/javascript",
            "application/json",
            "application/xml",
            "text/xml",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ]
        for doc in self:
            doc.is_text_editable = doc.mimetype in editable_types

    @api.depends("datas", "mimetype")
    def _compute_text_content(self):
        """Extract text content from document"""
        for doc in self:
            if not doc.datas:
                doc.text_content = ""
                continue

            try:
                decoded = base64.b64decode(doc.datas)

                # Handle DOCX files
                if (
                    doc.mimetype
                    == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ):
                    try:
                        import docx
                        from io import BytesIO

                        docx_file = docx.Document(BytesIO(decoded))
                        text = "\n".join([para.text for para in docx_file.paragraphs])
                        doc.text_content = text
                    except ImportError:
                        # Fallback: try mammoth for better conversion
                        try:
                            import mammoth
                            from io import BytesIO

                            result = mammoth.extract_raw_text(BytesIO(decoded))
                            doc.text_content = result.value
                        except ImportError:
                            doc.text_content = (
                                "[DOCX files require python-docx or mammoth library]"
                            )
                    except Exception as e:
                        _logger.error(f"Error reading DOCX: {e}")
                        doc.text_content = f"[Error reading DOCX: {str(e)}]"

                # Handle text files
                else:
                    doc.text_content = decoded.decode("utf-8", errors="replace")

            except Exception as e:
                _logger.error(f"Error decoding document: {e}")
                doc.text_content = f"[Error: {str(e)}]"

    def _inverse_text_content(self):
        """Save text content back to document"""
        for doc in self:
            if doc.text_content:
                try:
                    # For DOCX, save as plain text (you could enhance this to maintain formatting)
                    if (
                        doc.mimetype
                        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ):
                        # Convert back to DOCX
                        try:
                            import docx
                            from io import BytesIO

                            document = docx.Document()
                            for line in doc.text_content.split("\n"):
                                document.add_paragraph(line)

                            docx_bytes = BytesIO()
                            document.save(docx_bytes)
                            docx_bytes.seek(0)
                            doc.datas = base64.b64encode(docx_bytes.read())
                        except ImportError:
                            # Fallback: save as .txt
                            encoded = doc.text_content.encode("utf-8")
                            doc.datas = base64.b64encode(encoded)
                            doc.mimetype = "text/plain"
                            if not doc.name.endswith(".txt"):
                                doc.name = doc.name.rsplit(".", 1)[0] + ".txt"
                    else:
                        # Save as text
                        encoded = doc.text_content.encode("utf-8")
                        doc.datas = base64.b64encode(encoded)

                except Exception as e:
                    _logger.error(f"Error saving document: {e}")
                    raise UserError(f"Error saving document: {str(e)}")

    def action_open_editor(self):
        """Open document in custom editor"""
        self.ensure_one()
        return {
            "type": "ir.actions.client",
            "tag": "odoo_text_editor",
            "target": "new",
            "params": {
                "document_id": self.id,
                "document_name": self.name,
            },
        }
